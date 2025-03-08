from docx import Document
import pandas as pd
import tkinter as tk
import customtkinter as ctk
from tkinter import filedialog
from CTkMessagebox import CTkMessagebox

def limpezaDeDados():
    
    filename = filedialog.askopenfilename(
    title="Selecione um arquivo CSV",
    filetypes=[("Arquivos CSV", "*.csv")]
    )
    
    caminhoDoArquivo = filename
    # elimina os traços da coluna Solicitacao e ordena os dados por Razao social
    try:
        df = pd.read_csv(f"{caminhoDoArquivo}", sep=";", encoding="utf-8").sort_values(by="Razão social", ascending=True)
        df["Solicitação"] = df["Solicitação"].str.replace("-", " ", regex=True) #Remove traços da coluna Solicitação, substituindo por espaços 
        df.to_csv("consultaOS_modificado.csv", index=False) #salva arquivo csv
    
        labelAviso.configure(text="Limpeza concluida com sucesso!", text_color='green')
        labelAviso.pack()
         # Remove a label após 3 segundos (3000 ms)
        labelAviso.after(3000, lambda: labelAviso.pack_forget())
        emitirAviso()
    except:
        labelAviso.configure(text="Erro ao efetuar limpeza", text_color='red')
        labelAviso.pack()
         # Remove a label após 3 segundos (3000 ms)
        labelAviso.after(3000, lambda: labelAviso.pack_forget())

def emitirAviso():

    try:
        response = CTkMessagebox(title="Aviso", message="Selecione a pasta onde deseja salvar o arquivo!", icon="info", option_1="OK").get()
        if  response == "OK":
            exportarDoc()
            labelAviso.configure(text="Documento gerado com sucesso!", text_color='green')
            labelAviso.pack()
            labelAviso.after(3000, lambda: labelAviso.pack_forget())
        
    except:
        labelAviso.configure(text="Não foi possível gerar o documento!", text_color='red')
        labelAviso.pack()
        labelAviso.after(3000, lambda: labelAviso.pack_forget())
        

def exportarDoc():
    doc = Document()
    df = pd.read_csv("consultaOS_modificado.csv", encoding="utf-8") # lendo o arquivo csv
    
    pastaDestino = filedialog.asksaveasfilename(
        title="Salvar arquivo como...",
        defaultextension=".docx",
        filetypes=[("Arquivo docx", "*.docx")]
    )

    for index, row in df.iterrows(): 
        doc.add_paragraph(f"Razão social {row['Razão social']}")
        doc.add_paragraph(f"Solicitante: {row['Solicitante']}")
        doc.add_paragraph(f"Módulo: {row['Módulo']}")
        doc.add_paragraph(f"DLL:  {row['DLL']}")
        p = doc.add_paragraph()
        p.add_run('OS: ').bold = True
        p.add_run(f"{row['Numero']}")
        s = doc.add_paragraph()
        s.add_run('Solicitação: ').bold = True
        s.add_run(f"{row['Solicitação']}")
        doc.add_paragraph("\n" + "-" * 30 + "\n")  # Separador entre registros
    doc.save(pastaDestino) #salvar arquivo em pasta especificada por usuário

   

app = ctk.CTk()
app.title('Aviso de OS')
app.iconbitmap('img//definicoes.ico')
app.geometry('450x200')
app.resizable(False, False)


# Título da tela
tituloTela = ctk.CTkLabel(app, text="Aviso de OS", font=('Arial', 25))
tituloTela.pack(pady=20, padx=20, fill="x")

# lbl de avisos e críticas
labelAviso = ctk.CTkLabel(app, text=" ", font=('Arial',18))
labelAviso.pack()

# Frame para os campos CSV (lado esquerdo)
frameCSV = ctk.CTkFrame(app, width=450)
frameCSV.place(x=0, y=100)  # Posicionado com place para controle absoluto

labelCaminho = ctk.CTkLabel(frameCSV, text='Caminho do CSV:', font=('Arial', 20), anchor='nw')
labelCaminho.pack(side="left", padx=10, pady=10)
campoCaminhoCSV = ctk.CTkButton(frameCSV, command=limpezaDeDados,text="Selecione o arquivo CSV", width=270)
campoCaminhoCSV.pack(side="left", padx=0, pady=0)




app.mainloop()

